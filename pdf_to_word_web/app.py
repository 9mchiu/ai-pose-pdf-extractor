#!/usr/bin/env python3
"""
PDF → Word 轉換網頁工具
========================

使用者把 PDF 拖進網頁 → 自動萃取內容 → 下載 Word (.docx)。
- 文字部分：以文字呈現（掃描版也能讀，Claude 視覺 OCR）。
- 圖片 / 圖表部分：以圖片呈現（從原頁裁切後嵌入 Word）。

做法：
    每頁渲染成高解析圖 → 請 Claude 依閱讀順序回傳「區塊」
    (文字區塊給文字；圖片區塊給標註框 bbox) → 後端組成 .docx。

啟動：
    pip install -r requirements.txt
    export ANTHROPIC_API_KEY="你的 API key"
    python app.py
    然後打開瀏覽器: http://127.0.0.1:5000
"""

import io
import os

import fitz  # PyMuPDF
import anthropic
from PIL import Image
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# ────────────────────────────── 設定 ──────────────────────────────
MODEL = "claude-opus-4-8"     # 支援高解析視覺 + bbox 定位
RENDER_DPI = 200              # 頁面渲染解析度：越高圖越清楚但越慢/越貴
MAX_IMG_WIDTH_IN = 6.0        # Word 內嵌圖片最大寬度 (英吋)

# Claude 回傳的頁面結構 (JSON Schema)
PAGE_SCHEMA = {
    "type": "object",
    "properties": {
        "blocks": {
            "type": "array",
            "description": "此頁的內容區塊，依由上到下、由左到右的閱讀順序排列",
            "items": {
                "type": "object",
                "properties": {
                    "type": {"type": "string", "enum": ["text", "figure"]},
                    "content": {
                        "type": "string",
                        "description": "type=text 時放該區塊的完整文字(可含換行)；type=figure 時填空字串",
                    },
                    "bbox": {
                        "type": "array",
                        "items": {"type": "number"},
                        "description": "type=figure 時填 [x0,y0,x1,y1]，皆為 0~1 之間、相對整頁的比例；type=text 時填空陣列",
                    },
                    "caption": {
                        "type": "string",
                        "description": "figure 的圖說(若有)；否則填空字串",
                    },
                },
                "required": ["type", "content", "bbox", "caption"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["blocks"],
    "additionalProperties": False,
}

PROMPT = """你在把一頁 PDF（可能是掃描件）轉換成 Word 文件。

請依「閱讀順序」把這一頁拆成一連串區塊 (blocks)：
- 純文字的段落 / 標題 / 表格文字 → type="text"，content 放完整文字（保留自然換行）。
- 照片、示意圖、圖表、流程圖等視覺內容 → type="figure"，並給出它在整頁中的位置框 bbox=[x0,y0,x1,y1]。
  bbox 座標是 0~1 的比例（x 由左到右、y 由上到下），框請稍微寬鬆一點、把整張圖含圖說框進去。
  figure 的 content 留空字串，caption 放圖說文字（若有）。

規則：
- 只轉錄實際看得到的文字，不要翻譯、不要臆測、不要補內容。
- 文字要逐字忠實轉錄；數字、單位、符號都照原樣。
- 區塊順序要符合人閱讀的先後。
- 若整頁就是一張圖（例如全頁照片），就回一個 figure 區塊、bbox 約為 [0,0,1,1]。

依指定 JSON 結構輸出。"""

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB 上限


def analyze_page(client, png_bytes):
    """把一頁的 PNG 交給 Claude，回傳 blocks 清單。"""
    import base64
    b64 = base64.standard_b64encode(png_bytes).decode()
    resp = client.messages.create(
        model=MODEL,
        max_tokens=16000,
        output_config={"format": {"type": "json_schema", "schema": PAGE_SCHEMA}},
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": PROMPT},
                {"type": "image",
                 "source": {"type": "base64", "media_type": "image/png", "data": b64}},
            ],
        }],
    )
    import json
    text = next((b.text for b in resp.content if b.type == "text"), "{}")
    return json.loads(text).get("blocks", [])


def crop_figure(page_img: Image.Image, bbox):
    """依 0~1 比例的 bbox 從整頁圖裁切出圖片。"""
    w, h = page_img.size
    x0, y0, x1, y1 = bbox
    # 夾在合理範圍內，避免壞座標
    x0, x1 = sorted((max(0.0, min(1.0, x0)), max(0.0, min(1.0, x1))))
    y0, y1 = sorted((max(0.0, min(1.0, y0)), max(0.0, min(1.0, y1))))
    box = (int(x0 * w), int(y0 * h), int(x1 * w), int(y1 * h))
    if box[2] - box[0] < 5 or box[3] - box[1] < 5:      # 太小視為無效
        return None
    return page_img.crop(box)


def add_figure(doc: Document, crop: Image.Image, caption: str):
    buf = io.BytesIO()
    crop.save(buf, format="PNG")
    buf.seek(0)
    # 依比例決定寬度，不超過最大寬
    width_in = min(MAX_IMG_WIDTH_IN, crop.size[0] / RENDER_DPI)
    doc.add_picture(buf, width=Inches(max(1.0, width_in)))
    if caption and caption.strip():
        p = doc.add_paragraph()
        run = p.add_run(caption.strip())
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_docx(pdf_bytes, client) -> io.BytesIO:
    doc = Document()
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    mat = fitz.Matrix(RENDER_DPI / 72, RENDER_DPI / 72)

    for i, page in enumerate(src):
        png_bytes = page.get_pixmap(matrix=mat).tobytes("png")
        page_img = Image.open(io.BytesIO(png_bytes))

        try:
            blocks = analyze_page(client, png_bytes)
        except Exception as e:
            # 該頁失敗 → 直接把整頁當一張圖放進去，不中斷
            add_figure(doc, page_img, f"[第 {i+1} 頁自動判讀失敗，改以整頁圖片呈現]")
            doc.add_page_break()
            continue

        for blk in blocks:
            if blk.get("type") == "figure":
                crop = crop_figure(page_img, blk.get("bbox") or [0, 0, 1, 1])
                if crop is not None:
                    add_figure(doc, crop, blk.get("caption", ""))
            else:
                content = (blk.get("content") or "").strip()
                for line in content.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

        if i < len(src) - 1:
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/convert", methods=["POST"])
def convert():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        return jsonify(error="伺服器未設定 ANTHROPIC_API_KEY"), 500
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".pdf"):
        return jsonify(error="請上傳 PDF 檔"), 400

    try:
        client = anthropic.Anthropic()
        docx_io = build_docx(f.read(), client)
    except Exception as e:
        return jsonify(error=f"轉換失敗：{e}"), 500

    out_name = os.path.splitext(os.path.basename(f.filename))[0] + ".docx"
    return send_file(
        docx_io,
        as_attachment=True,
        download_name=out_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    print("PDF → Word 轉換器已啟動： http://127.0.0.1:5000")
    app.run(host="127.0.0.1", port=5000, debug=False)
